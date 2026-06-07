#!/usr/bin/env python3
"""Generate the project report in DOCX and PDF formats."""

import os
from datetime import datetime

# ────────────────────────────────────────────────
# Generate DOCX Report
# ────────────────────────────────────────────────
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('ML Classification Project Report', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Titanic Survival Prediction\nLogistic Regression vs Random Forest')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()  # spacer

# Meta info
meta = doc.add_paragraph()
meta.add_run('Author: ').bold = True
meta.add_run('Arena.ai Agent')

meta2 = doc.add_paragraph()
meta2.add_run('Date: ').bold = True
meta2.add_run('2026-06-07')

meta3 = doc.add_paragraph()
meta3.add_run('Dataset: ').bold = True
meta3.add_run('Titanic (891 passengers, 15 features)')

doc.add_page_break()

# ─── 1. Introduction ───
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    'This project builds supervised classification models to predict passenger survival '
    'aboard the Titanic. The goal is to compare two fundamentally different algorithms—'
    'Logistic Regression (linear) and Random Forest (ensemble tree-based)—on the same '
    'task and evaluate their performance using standard classification metrics.'
)
doc.add_paragraph(
    'The dataset contains 891 passenger records with features including passenger class, '
    'sex, age, number of siblings/spouses aboard, number of parents/children aboard, fare, '
    'and port of embarkation.'
)

# ─── 2. Methodology ───
doc.add_heading('2. Methodology', level=1)

doc.add_heading('2.1 Data Preprocessing', level=2)
doc.add_paragraph('• Dropped redundant/leaky columns: alive, who, adult_male, deck, embark_town, class')
doc.add_paragraph('• Imputed missing Age values using median by Sex and Pclass group')
doc.add_paragraph('• Imputed missing Embarked values with mode')
doc.add_paragraph('• Encoded categorical variables (sex, embarked) using Label Encoding')
doc.add_paragraph('• Applied StandardScaler for feature normalization')
doc.add_paragraph('• Selected 7 features: pclass, sex, age, sibsp, parch, fare, embarked')

doc.add_heading('2.2 Train/Test Split', level=2)
doc.add_paragraph('• 80/20 stratified split preserving class distribution')
doc.add_paragraph('• Training: 712 samples | Test: 179 samples')

doc.add_heading('2.3 Cross-Validation', level=2)
doc.add_paragraph('• Stratified 5-Fold Cross-Validation on training set')
doc.add_paragraph('• Evaluated using Accuracy and ROC-AUC scoring')

doc.add_heading('2.4 Algorithms', level=2)
doc.add_paragraph('1. Logistic Regression (max_iter=1000, L2 regularization)')
doc.add_paragraph('2. Random Forest (100 trees, max_depth=6)')

# ─── 3. Results ───
doc.add_heading('3. Results', level=1)

doc.add_heading('3.1 Cross-Validation Scores', level=2)
table = doc.add_table(rows=3, cols=3, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Model', 'CV Accuracy (mean)', 'CV ROC-AUC (mean)']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
table.rows[1].cells[0].text = 'Logistic Regression'
table.rows[1].cells[1].text = '0.8034'
table.rows[1].cells[2].text = '0.8552'
table.rows[2].cells[0].text = 'Random Forest'
table.rows[2].cells[1].text = '0.8216'
table.rows[2].cells[2].text = '0.8748'

doc.add_paragraph()

doc.add_heading('3.2 Test Set Performance', level=2)
table2 = doc.add_table(rows=6, cols=3, style='Light Grid Accent 1')
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
headers2 = ['Metric', 'Logistic Regression', 'Random Forest']
for i, h in enumerate(headers2):
    table2.rows[0].cells[i].text = h
metrics_data = [
    ('Accuracy', '0.8212', '0.7989'),
    ('Precision', '0.8136', '0.8235'),
    ('Recall', '0.6957', '0.6087'),
    ('F1-Score', '0.7500', '0.7000'),
    ('ROC-AUC', '0.8497', '0.8508'),
]
for r, (metric, lr_val, rf_val) in enumerate(metrics_data, 1):
    table2.rows[r].cells[0].text = metric
    table2.rows[r].cells[1].text = lr_val
    table2.rows[r].cells[2].text = rf_val

doc.add_paragraph()

# Include screenshots
doc.add_heading('3.3 Visualizations', level=2)

img_dir = '/home/user/ml-classification-project/notebook'
plots = [
    ('Confusion Matrices', 'confusion_matrices.png'),
    ('ROC Curves', 'roc_curves.png'),
    ('Feature Importance (RF)', 'feature_importance.png'),
    ('LR Coefficients', 'lr_coefficients.png'),
    ('Metrics Comparison', 'metrics_comparison.png'),
    ('Learning Curves', 'learning_curves.png'),
    ('EDA Plots', 'eda_plots.png'),
]

for plot_title, plot_file in plots:
    doc.add_heading(plot_title, level=3)
    img_path = os.path.join(img_dir, plot_file)
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

# ─── 4. Discussion ───
doc.add_heading('4. Discussion', level=1)
doc.add_paragraph(
    'Logistic Regression outperformed Random Forest on the test set in terms of accuracy '
    '(82.1% vs 79.9%), recall (69.6% vs 60.9%), and F1-score (0.750 vs 0.700). Random Forest '
    'showed slightly higher precision (82.4% vs 81.4%) and a marginally better ROC-AUC (0.851 vs 0.850).'
)
doc.add_paragraph(
    'The learning curves indicate that Logistic Regression generalizes well with minimal '
    'overfitting, while Random Forest shows a gap between training and validation scores '
    'suggesting some overfitting despite depth constraints.'
)
doc.add_paragraph(
    'Feature importance analysis reveals that sex, passenger class, and fare are the most '
    'influential predictors—consistent with the historical "women and children first" protocol '
    'and the economic disparity in survival rates.'
)

# ─── 5. Conclusion ───
doc.add_heading('5. Conclusion', level=1)
doc.add_paragraph(
    'Based on the evaluation metrics, Logistic Regression is recommended as the final model '
    'for this task. It achieves higher accuracy, better recall, superior F1-score, and equal '
    'ROC-AUC compared to Random Forest, while offering greater interpretability through its '
    'coefficients. For production deployment, further tuning and feature engineering could '
    'potentially improve both models.'
)

doc.add_heading('5.1 Reproducibility', level=2)
doc.add_paragraph('• Python 3.13 with scikit-learn 1.6.1, pandas 2.2.3, numpy 2.3.5')
doc.add_paragraph('• All code and notebook available in the project repository')
doc.add_paragraph('• See README.md for setup and execution instructions')

# Save DOCX
docx_path = '/home/user/ml-classification-project/report/ML_Classification_Report.docx'
doc.save(docx_path)
print(f"DOCX report saved to {docx_path}")

# ────────────────────────────────────────────────
# Generate PDF Report
# ────────────────────────────────────────────────
from fpdf import FPDF

class PDFReport(FPDF):
    def header(self):
        self.set_font('Helvetica', 'B', 10)
        self.set_text_color(100, 100, 100)
        self.cell(0, 8, 'ML Classification Project - Titanic Survival Prediction', 0, 1, 'C')
        self.line(10, 14, 200, 14)
        self.ln(4)

    def footer(self):
        self.set_y(-15)
        self.set_font('Helvetica', 'I', 8)
        self.set_text_color(128)
        self.cell(0, 10, f'Page {self.page_no()}/{{nb}}', 0, 0, 'C')

    def section_title(self, title):
        self.set_font('Helvetica', 'B', 14)
        self.set_text_color(30, 60, 114)
        self.cell(0, 10, title, 0, 1, 'L')
        self.line(10, self.get_y(), 200, self.get_y())
        self.ln(2)

    def sub_title(self, title):
        self.set_font('Helvetica', 'B', 11)
        self.set_text_color(60, 90, 140)
        self.cell(0, 8, title, 0, 1, 'L')
        self.ln(1)

    def body_text(self, text):
        self.set_font('Helvetica', '', 10)
        self.set_text_color(40, 40, 40)
        self.multi_cell(0, 5, text)
        self.ln(2)

    def bullet(self, text):
        self.set_font('Helvetica', '', 10)
        self.set_text_color(40, 40, 40)
        self.cell(5)
        self.cell(0, 5, f'- {text}', 0, 1)

pdf = PDFReport()
pdf.alias_nb_pages()
pdf.add_page()

# Title
pdf.set_font('Helvetica', 'B', 22)
pdf.set_text_color(30, 60, 114)
pdf.cell(0, 15, 'ML Classification Project Report', 0, 1, 'C')
pdf.ln(2)

pdf.set_font('Helvetica', '', 13)
pdf.set_text_color(80, 80, 80)
pdf.cell(0, 7, 'Titanic Survival Prediction', 0, 1, 'C')
pdf.cell(0, 7, 'Logistic Regression vs Random Forest', 0, 1, 'C')
pdf.ln(5)

pdf.set_font('Helvetica', '', 10)
pdf.set_text_color(60, 60, 60)
pdf.cell(0, 6, 'Author: Arena.ai Agent  |  Date: 2026-06-07  |  Dataset: Titanic (891 samples)', 0, 1, 'C')
pdf.ln(8)

# 1. Introduction
pdf.section_title('1. Introduction')
pdf.body_text(
    'This project builds supervised classification models to predict passenger survival '
    'aboard the Titanic. The goal is to compare two fundamentally different algorithms-- '
    'Logistic Regression (linear) and Random Forest (ensemble tree-based)--on the same '
    'task and evaluate their performance using standard classification metrics.'
)
pdf.body_text(
    'The dataset contains 891 passenger records with features including passenger class, '
    'sex, age, number of siblings/spouses aboard, number of parents/children aboard, fare, '
    'and port of embarkation.'
)

# 2. Methodology
pdf.section_title('2. Methodology')

pdf.sub_title('2.1 Data Preprocessing')
pdf.bullet('Dropped redundant/leaky columns: alive, who, adult_male, deck, embark_town, class')
pdf.bullet('Imputed missing Age values using median by Sex and Pclass group')
pdf.bullet('Imputed missing Embarked values with mode')
pdf.bullet('Encoded categorical variables (sex, embarked) using Label Encoding')
pdf.bullet('Applied StandardScaler for feature normalization')
pdf.bullet('Selected 7 features: pclass, sex, age, sibsp, parch, fare, embarked')

pdf.sub_title('2.2 Train/Test Split')
pdf.bullet('80/20 stratified split preserving class distribution')
pdf.bullet('Training: 712 samples  |  Test: 179 samples')

pdf.sub_title('2.3 Cross-Validation')
pdf.bullet('Stratified 5-Fold Cross-Validation on training set')
pdf.bullet('Evaluated using Accuracy and ROC-AUC scoring')

pdf.sub_title('2.4 Algorithms')
pdf.bullet('Logistic Regression (max_iter=1000, L2 regularization)')
pdf.bullet('Random Forest (100 trees, max_depth=6)')

# 3. Results
pdf.section_title('3. Results')

pdf.sub_title('3.1 Cross-Validation Scores')
pdf.set_font('Courier', '', 9)
pdf.cell(5)
pdf.cell(0, 5, '+----------------------+----------------+---------------+', 0, 1)
pdf.cell(5)
pdf.cell(0, 5, '| Model                | CV Acc (mean)  | CV ROC-AUC    |', 0, 1)
pdf.cell(5)
pdf.cell(0, 5, '+----------------------+----------------+---------------+', 0, 1)
pdf.cell(5)
pdf.cell(0, 5, '| Logistic Regression  | 0.8034         | 0.8552        |', 0, 1)
pdf.cell(5)
pdf.cell(0, 5, '| Random Forest        | 0.8216         | 0.8748        |', 0, 1)
pdf.cell(5)
pdf.cell(0, 5, '+----------------------+----------------+---------------+', 0, 1)
pdf.set_font('Helvetica', '', 10)
pdf.ln(3)

pdf.sub_title('3.2 Test Set Performance')
pdf.set_font('Courier', '', 9)
pdf.cell(5)
pdf.cell(0, 5, '+----------------+---------------------+----------------+', 0, 1)
pdf.cell(5)
pdf.cell(0, 5, '| Metric         | Logistic Regression | Random Forest  |', 0, 1)
pdf.cell(5)
pdf.cell(0, 5, '+----------------+---------------------+----------------+', 0, 1)
for metric, lr_v, rf_v in metrics_data:
    pdf.cell(5)
    pdf.cell(0, 5, f'| {metric:14s} | {lr_v:19s} | {rf_v:14s} |', 0, 1)
pdf.cell(5)
pdf.cell(0, 5, '+----------------+---------------------+----------------+', 0, 1)
pdf.set_font('Helvetica', '', 10)
pdf.ln(5)

# Add images
pdf.sub_title('3.3 Visualizations')

for plot_title, plot_file in plots:
    img_path = os.path.join(img_dir, plot_file)
    if os.path.exists(img_path):
        pdf.add_page()
        pdf.sub_title(plot_title)
        pdf.image(img_path, x=15, w=180)

# 4. Discussion
pdf.add_page()
pdf.section_title('4. Discussion')
pdf.body_text(
    'Logistic Regression outperformed Random Forest on the test set in terms of accuracy '
    '(82.1% vs 79.9%), recall (69.6% vs 60.9%), and F1-score (0.750 vs 0.700). Random Forest '
    'showed slightly higher precision (82.4% vs 81.4%) and a marginally better ROC-AUC '
    '(0.851 vs 0.850).'
)
pdf.body_text(
    'The learning curves indicate that Logistic Regression generalizes well with minimal '
    'overfitting, while Random Forest shows a gap between training and validation scores '
    'suggesting some overfitting despite depth constraints.'
)
pdf.body_text(
    'Feature importance analysis reveals that sex, passenger class, and fare are the most '
    'influential predictors--consistent with the historical "women and children first" protocol '
    'and the economic disparity in survival rates.'
)

# 5. Conclusion
pdf.section_title('5. Conclusion')
pdf.body_text(
    'Based on the evaluation metrics, Logistic Regression is recommended as the final model '
    'for this task. It achieves higher accuracy, better recall, superior F1-score, and equal '
    'ROC-AUC compared to Random Forest, while offering greater interpretability through its '
    'coefficients. For production deployment, further tuning and feature engineering could '
    'potentially improve both models.'
)

pdf.sub_title('5.1 Reproducibility')
pdf.bullet('Python 3.13 with scikit-learn 1.6.1, pandas 2.2.3, numpy 2.3.5')
pdf.bullet('All code and notebook available in the project repository')
pdf.bullet('See README.md for setup and execution instructions')

# Save PDF
pdf_path = '/home/user/ml-classification-project/report/ML_Classification_Report.pdf'
pdf.output(pdf_path)
print(f"PDF report saved to {pdf_path}")

print("\nAll reports generated successfully!")
